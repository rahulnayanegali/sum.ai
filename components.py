import json
from docx import Document
import os
import speech_recognition as sr
import time
from datetime import datetime
from openai import OpenAI

def record_text(retries=3):
    # Clear the output.txt file at the start of recording
    with open("output.txt", "w") as f:
        f.write("")

    r = sr.Recognizer()
    for attempt in range(retries):
        try:
            with sr.Microphone() as source2:
                r.adjust_for_ambient_noise(source2, duration=0.5)
                print("Please say something")
                audio2 = r.listen(source2, timeout=5, phrase_time_limit=40)
                MyText = r.recognize_google(audio2)
                return MyText
        except sr.WaitTimeoutError:
            print("Listening timed out while waiting for phrase to start. Retrying...")
        except sr.RequestError as e:
            print(f"Could not request results; {e}. Retrying...")
            time.sleep(1)
        except sr.UnknownValueError:
            print("Unknown error occurred. Retrying...")
            time.sleep(1)
        except ConnectionResetError as e:
            print(f"Connection error: {e}. Retrying...")
            time.sleep(1)
    print("Failed to recognize speech after multiple attempts.")
    return None

def output_text_buffer(text_buffer, file_path="output.txt"):
    with open(file_path, "a") as f:
        for text in text_buffer:
            f.write(text + "\n")

def read_text_file(file_path):
    with open(file_path, "r") as file:
        content = file.read()
    return content

def load_gpt_key():
    with open("config.json") as f:
        config_data = json.load(f)
    return config_data["GPT_KEY"]

def summarize_text(text, gpt_key):
    client = OpenAI(api_key=gpt_key)

    prompt = f"Write a summary with the subheading and the discussion about the team below it:\n\n{text}"
    
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=150,
        temperature=0.7
    )
    
    summary = response.choices[0].message.content.strip()
    return summary

def output_summary_to_docx(summary, file_path="summary.docx"):
    doc = Document()
    today = datetime.today().strftime('%Y-%m-%d')
    doc.add_heading(f'Meeting Summary - {today}', 0)

    # Split summary into lines and format subheadings
    lines = summary.split('\n')
    for line in lines:
        if line.startswith("**") and line.endswith("**"):
            subheading = line[2:-2].strip()  # Remove the ** and strip any extra whitespace
            p = doc.add_paragraph()
            run = p.add_run(subheading)
            run.bold = True
            run.underline = True
            run.font.size = doc.styles['Heading 2'].font.size  # Set to the size of Heading 2
        elif line.startswith("**"):
            subheading = line[2:].strip()  # Remove the ** at the beginning and strip any extra whitespace
            p = doc.add_paragraph()
            run = p.add_run(subheading)
            run.bold = True
            run.underline = True
            run.font.size = doc.styles['Heading 2'].font.size  # Set to the size of Heading 2
        else:
            doc.add_paragraph(line)

    doc.save(file_path)
